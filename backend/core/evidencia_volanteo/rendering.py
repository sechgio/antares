"""Renderizado de Evidencia Volanteo a PDF (WeasyPrint) y DOCX (python-docx)."""

from __future__ import annotations

import base64
import contextlib
import logging
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import TYPE_CHECKING, Any

from jinja2 import Environment, FileSystemLoader, select_autoescape

from backend.utils.pdf_html import write_pdf_sanitized

from .errors import RenderingError
from .layout import (
    BORDER_PT,
    CUADRANTE_LABEL,
    EMPTY_CUADRANTE_PLACEHOLDER,
    GAP_UNDER_HEADER_CM,
    HEADER_INFO_HEIGHT_CM,
    HEADER_LOGO_WIDTH_CM,
    HEADER_TITLE_HEIGHT_CM,
    HEADER_TITLE_WIDTH_CM,
    INFO_FONT_PT,
    LOGO_MAX_HEIGHT_CM,
    LOGO_MAX_WIDTH_CM,
    PHOTO_COLS,
    PHOTO_GAP_CM,
    PHOTO_HEIGHT_CM,
    PHOTO_ROWS,
    PHOTO_TABLE_COLS,
    PHOTO_TABLE_ROWS,
    PHOTO_WIDTH_CM,
    TABLE_WIDTH_CM,
    TITLE_FONT_PT,
    layout_context,
)
from .models import MAX_PAGES

if TYPE_CHECKING:
    from .models import EvidenciaDocument

logger = logging.getLogger(__name__)

DOC_FONT = "Aptos"
MAX_SLOTS = PHOTO_COLS * PHOTO_ROWS
BORDER_SZ = str(int(BORDER_PT * 8))


def _display_cuadrante(value: str) -> str:
    stripped = value.strip()
    if not stripped:
        return EMPTY_CUADRANTE_PLACEHOLDER
    return stripped.upper()


def _breakable_text(value: str, chunk: int = 20) -> str:
    """Inserta saltos opcionales para que Word envuelva tokens largos sin espacios."""
    if len(value) <= chunk or any(ch.isspace() for ch in value):
        return value
    zwsp = "\u200b"
    return zwsp.join(value[i : i + chunk] for i in range(0, len(value), chunk))


def _table_borders_xml() -> str:
    return (
        '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:top w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="000000"/>'
        f'<w:left w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="000000"/>'
        f'<w:bottom w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="000000"/>'
        f'<w:right w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="000000"/>'
        f'<w:insideH w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="000000"/>'
        f'<w:insideV w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )


def _photos_table_borders_xml() -> str:
    """Solo marco exterior del panel; sin líneas internas."""
    return (
        '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:top w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="000000"/>'
        f'<w:left w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="000000"/>'
        f'<w:bottom w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="000000"/>'
        f'<w:right w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="nil"/>'
        '<w:insideV w:val="nil"/>'
        '</w:tblBorders>'
    )


def _nil_cell_borders_xml() -> str:
    return (
        '<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="nil"/>'
        '<w:bottom w:val="nil"/>'
        '<w:left w:val="nil"/>'
        '<w:right w:val="nil"/>'
        '</w:tcBorders>'
    )


def _cell_borders_xml(
    *,
    top: bool = False,
    bottom: bool = False,
    left: bool = False,
    right: bool = False,
) -> str:
    """Bordes de celda. En Word tcBorders anula tblBorders: el marco exterior
    del panel de fotos debe pintarse en las celdas del perímetro."""
    def side(name: str, on: bool) -> str:
        if on:
            return (
                f'<w:{name} w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="000000"/>'
            )
        return f'<w:{name} w:val="nil"/>'

    return (
        '<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'{side("top", top)}{side("left", left)}{side("bottom", bottom)}{side("right", right)}'
        '</w:tcBorders>'
    )


def _nil_table_borders_xml() -> str:
    return (
        '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="nil"/>'
        '<w:left w:val="nil"/>'
        '<w:bottom w:val="nil"/>'
        '<w:right w:val="nil"/>'
        '<w:insideH w:val="nil"/>'
        '<w:insideV w:val="nil"/>'
        '</w:tblBorders>'
    )


def _resolve_template_dir() -> Path:
    bundled = Path(__file__).resolve().parent.parent.parent / "templates"
    if bundled.exists():
        return bundled
    return Path(__file__).resolve().parent.parent / "templates"


_TEMPLATE_DIR = _resolve_template_dir()
_jinja_env = Environment(
    loader=FileSystemLoader(str(_TEMPLATE_DIR)),
    autoescape=select_autoescape(["html", "xml"]),
)


def _data_uri_from_b64(b64_string: str) -> str:
    if b64_string.startswith("data:"):
        header_end = b64_string.find(",")
        if header_end != -1:
            b64_string = b64_string[header_end + 1 :]
    mime = "image/png"
    try:
        sample = b64_string[:24] + "=" * ((4 - len(b64_string[:24]) % 4) % 4)
        header = base64.b64decode(sample, validate=True)
        if header.startswith(b"\xff\xd8"):
            mime = "image/jpeg"
        elif header.startswith(b"\x89PNG"):
            mime = "image/png"
        elif header.startswith(b"RIFF") and header[8:12] == b"WEBP":
            mime = "image/webp"
    except Exception:
        pass
    return f"data:{mime};base64,{b64_string}"


def _valid_image_bytes(content: bytes) -> bool:
    try:
        from PIL import Image

        with Image.open(BytesIO(content)) as image:
            image.verify()
        return True
    except Exception:
        return False


def _valid_b64_image(b64_string: str) -> bool:
    try:
        content = base64.b64decode(b64_string, validate=True)
    except Exception:
        return False
    return _valid_image_bytes(content)


def _contain_fit_cm(content: bytes, max_width_cm: float, max_height_cm: float) -> tuple[float, float]:
    try:
        from PIL import Image

        with Image.open(BytesIO(content)) as image:
            width_px, height_px = image.size
    except Exception:
        return max_width_cm, max_height_cm
    if width_px <= 0 or height_px <= 0:
        return max_width_cm, max_height_cm
    scale = min(max_width_cm / width_px, max_height_cm / height_px)
    return width_px * scale, height_px * scale


def _serialize_pages(
    document: EvidenciaDocument,
    image_uris: dict[str, str],
) -> list[dict[str, Any]]:
    pages_data: list[dict[str, Any]] = []
    for page in document.pages:
        slots: list[dict[str, Any] | None] = [None] * MAX_SLOTS
        for ref in page.images:
            if 1 <= ref.position <= MAX_SLOTS:
                slots[ref.position - 1] = {
                    "filename": ref.filename,
                    "uri": image_uris.get(ref.filename),
                }
        page_cuadrante = page.cuadrante or document.cuadrante
        pages_data.append({
            "slots": slots,
            "cuadrante": _display_cuadrante(page_cuadrante),
        })
    return pages_data


def _build_image_uris(
    images: dict[str, str],
    image_paths: dict[str, str] | None,
) -> dict[str, str]:
    image_uris: dict[str, str] = {}
    for filename, raw_path in (image_paths or {}).items():
        path = Path(raw_path)
        if path.is_file():
            with contextlib.suppress(Exception):
                if _valid_image_bytes(path.read_bytes()):
                    image_uris[filename] = path.resolve().as_uri()
    for filename, b64 in images.items():
        if filename in image_uris:
            continue
        if _valid_b64_image(b64):
            image_uris[filename] = _data_uri_from_b64(b64)
    return image_uris


def _prepare_logos(logos: dict[str, str | None]) -> tuple[str | None, str | None]:
    left_raw = logos.get("left")
    right_raw = logos.get("right")
    left = _data_uri_from_b64(left_raw) if left_raw and _valid_b64_image(left_raw) else None
    right = _data_uri_from_b64(right_raw) if right_raw and _valid_b64_image(right_raw) else None
    return left, right


def _default_filename(fmt: str) -> str:
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    ext = "docx" if fmt == "docx" else "pdf"
    return f"evidencia_volanteo_{ts}.{ext}"


def render_pdf_html(html_string: str) -> tuple[bytes, str]:
    if not html_string.strip():
        msg = "No hay HTML para exportar"
        raise RenderingError(msg)
    try:
        return write_pdf_sanitized(html_string), _default_filename("pdf")
    except RenderingError:
        raise
    except Exception as exc:
        logger.exception("Error al generar PDF desde HTML de evidencia volanteo")
        msg = f"Error al generar PDF: {exc}"
        raise RenderingError(msg) from exc


def render_pdf(
    document: EvidenciaDocument,
    logos: dict[str, str | None],
    images: dict[str, str],
    image_paths: dict[str, str] | None = None,
) -> tuple[bytes, str]:
    if not document.pages:
        msg = "No hay páginas para exportar"
        raise RenderingError(msg)
    if len(document.pages) > MAX_PAGES:
        msg = f"El PDF excede el máximo de {MAX_PAGES} páginas"
        raise RenderingError(msg)

    try:
        template = _jinja_env.get_template("evidencia-volanteo.html")
    except Exception as exc:
        msg = f"Error al cargar plantilla: {exc}"
        raise RenderingError(msg) from exc

    logo_left, logo_right = _prepare_logos(logos)
    image_uris = _build_image_uris(images, image_paths)
    pages_data = _serialize_pages(document, image_uris)
    filename = _default_filename("pdf")

    context = {
        "title": document.title,
        "cuadrante": document.cuadrante,
        "pages": pages_data,
        "logo_left": logo_left,
        "logo_right": logo_right,
        **layout_context(),
        "cuadrante_label": document.cuadrante_label or CUADRANTE_LABEL,
        "show_cuadrante_label": document.show_cuadrante_label,
    }

    try:
        html_string = template.render(context)
        # Images/logos are already data-URIs; do not open base_url file/http fetch.
        return write_pdf_sanitized(html_string), filename
    except RenderingError:
        raise
    except Exception as exc:
        logger.exception("Error al generar PDF de evidencia volanteo")
        msg = f"Error al generar PDF: {exc}"
        raise RenderingError(msg) from exc


def render_docx(
    document: EvidenciaDocument,
    logos: dict[str, str | None],
    images: dict[str, str],
    image_paths: dict[str, str] | None = None,
) -> tuple[bytes, str]:
    if not document.pages:
        msg = "No hay páginas para exportar"
        raise RenderingError(msg)
    if len(document.pages) > MAX_PAGES:
        msg = f"El documento excede el máximo de {MAX_PAGES} páginas"
        raise RenderingError(msg)

    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    def cm_to_twips(cm: float) -> int:
        return round(cm * 567)

    def set_cell_width(cell: Any, width_cm: float) -> None:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn("w:tcW")):
            tcPr.remove(old)
        tcPr.append(parse_xml(
            f'<w:tcW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            f'w:w="{cm_to_twips(width_cm)}" w:type="dxa"/>',
        ))

    def set_row_height(row: Any, height_cm: float, *, rule: str = "exact") -> None:
        """rule='exact' fija altura; 'atLeast' permite crecer (p. ej. cuadrante multilínea)."""
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        for old in trPr.findall(qn("w:trHeight")):
            trPr.remove(old)
        trPr.append(parse_xml(
            f'<w:trHeight xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            f'w:val="{cm_to_twips(height_cm)}" w:hRule="{rule}"/>',
        ))

    def set_vertical_align(cell: Any, align: str) -> None:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn("w:vAlign")):
            tcPr.remove(old)
        tcPr.append(parse_xml(
            f'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            f'w:val="{align}"/>',
        ))

    def format_run(run: Any, size_pt: float, *, bold: bool = False) -> None:
        run.bold = bold
        run.font.size = Pt(size_pt)
        run.font.name = DOC_FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)

    def reset_cell_paragraph(paragraph: Any, *, line_spacing: float = 1.0) -> None:
        """Neutraliza los defaults heredados (10pt after, 1.15x) que desplazan/recortan
        el contenido en celdas de altura exacta, replicando margin/padding 0 y
        line-height 0 del preview y del HTML."""
        pf = paragraph.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = line_spacing

    def pt_to_twips(pt: float) -> int:
        return round(pt * 20)

    def set_cell_margins(
        cell: Any,
        *,
        top_pt: float = 0,
        left_pt: float = 0,
        bottom_pt: float = 0,
        right_pt: float = 0,
    ) -> None:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn("w:tcMar")):
            tcPr.remove(old)
        tcPr.append(parse_xml(
            '<w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            f'<w:top w:w="{pt_to_twips(top_pt)}" w:type="dxa"/>'
            f'<w:left w:w="{pt_to_twips(left_pt)}" w:type="dxa"/>'
            f'<w:bottom w:w="{pt_to_twips(bottom_pt)}" w:type="dxa"/>'
            f'<w:right w:w="{pt_to_twips(right_pt)}" w:type="dxa"/>'
            '</w:tcMar>',
        ))

    def set_table_no_cell_margins(table: Any) -> None:
        """Anula los margenes de celda por defecto (108 twips L/R) para que las
        fotos llenen la celda exactamente, igual que padding 0 del preview/PDF."""
        tblPr = table._tbl.tblPr
        for old in tblPr.findall(qn("w:tblCellMar")):
            tblPr.remove(old)
        tblPr.append(parse_xml(
            '<w:tblCellMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:top w:w="0" w:type="dxa"/>'
            '<w:left w:w="0" w:type="dxa"/>'
            '<w:bottom w:w="0" w:type="dxa"/>'
            '<w:right w:w="0" w:type="dxa"/>'
            '</w:tblCellMar>',
        ))

    def fill_image_size_cm(content: bytes, max_w: float, max_h: float) -> tuple[float, float, tuple[int, int, int, int]]:
        return max_w, max_h, (0, 0, 0, 0)

    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(0.8)
    section.right_margin = Cm(0.8)

    image_bytes: dict[str, bytes] = {}
    for filename, b64 in images.items():
        with contextlib.suppress(Exception):
            image_bytes[filename] = base64.b64decode(b64, validate=True)
    for filename, raw_path in (image_paths or {}).items():
        path = Path(raw_path)
        if path.is_file() and filename not in image_bytes:
            with contextlib.suppress(Exception):
                image_bytes[filename] = path.read_bytes()

    logo_left_bytes: bytes | None = None
    logo_right_bytes: bytes | None = None
    left_b64 = logos.get("left")
    if left_b64:
        with contextlib.suppress(Exception):
            logo_left_bytes = base64.b64decode(left_b64, validate=True)
    right_b64 = logos.get("right")
    if right_b64:
        with contextlib.suppress(Exception):
            logo_right_bytes = base64.b64decode(right_b64, validate=True)

    logo_left_dims = (
        _contain_fit_cm(logo_left_bytes, LOGO_MAX_WIDTH_CM, LOGO_MAX_HEIGHT_CM)
        if logo_left_bytes else (0.0, 0.0)
    )
    logo_right_dims = (
        _contain_fit_cm(logo_right_bytes, LOGO_MAX_WIDTH_CM, LOGO_MAX_HEIGHT_CM)
        if logo_right_bytes else (0.0, 0.0)
    )

    for page_idx, page in enumerate(document.pages):
        # 1. Header Table: 2 rows x 3 cols
        header_table = doc.add_table(rows=2, cols=3)
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        header_table.autofit = False
        header_table.allow_autofit = False

        if page_idx > 0:
            first_cell = header_table.cell(0, 0)
            pPr = first_cell.paragraphs[0]._p.get_or_add_pPr()
            pPr.append(parse_xml(
                '<w:pageBreakBefore xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>',
            ))

        # Set borders for header_table
        tblPr = header_table._tbl.tblPr
        for tag in ("w:tblBorders",):
            existing = tblPr.find(qn(tag))
            if existing is not None:
                tblPr.remove(existing)
        tblPr.append(parse_xml(_table_borders_xml()))
        tblPr.append(parse_xml(
            '<w:tblW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            f'w:w="{cm_to_twips(TABLE_WIDTH_CM)}" w:type="dxa"/>',
        ))
        tblPr.append(parse_xml(
            '<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            'w:type="fixed"/>',
        ))
        set_table_no_cell_margins(header_table)

        # Set header column widths
        header_widths = [HEADER_LOGO_WIDTH_CM, HEADER_TITLE_WIDTH_CM, HEADER_LOGO_WIDTH_CM]
        grid = header_table._tbl.tblGrid
        for col, width_cm in zip(grid.gridCol_lst, header_widths, strict=True):
            col.set(qn("w:w"), str(cm_to_twips(width_cm)))

        # Header heights — info usa atLeast: exact recorta el cuadrante multilínea en Word
        set_row_height(header_table.rows[0], HEADER_TITLE_HEIGHT_CM)
        set_row_height(header_table.rows[1], HEADER_INFO_HEIGHT_CM, rule="atLeast")

        # Merge logo cells vertically
        header_table.cell(0, 0).merge(header_table.cell(1, 0))
        header_table.cell(0, 2).merge(header_table.cell(1, 2))

        # Populate Logo Left
        merged_left = header_table.cell(0, 0)
        set_cell_width(merged_left, HEADER_LOGO_WIDTH_CM)
        set_cell_margins(merged_left, top_pt=2, left_pt=4, bottom_pt=2, right_pt=4)
        set_vertical_align(merged_left, "center")
        merged_left.paragraphs[0].clear()
        reset_cell_paragraph(merged_left.paragraphs[0])
        merged_left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if logo_left_bytes:
            run = merged_left.paragraphs[0].add_run()
            run.add_picture(BytesIO(logo_left_bytes), width=Cm(logo_left_dims[0]), height=Cm(logo_left_dims[1]))

        # Populate Title (with underline!)
        title_cell = header_table.cell(0, 1)
        set_cell_width(title_cell, HEADER_TITLE_WIDTH_CM)
        set_cell_margins(title_cell, top_pt=2, left_pt=4, bottom_pt=2, right_pt=4)
        set_vertical_align(title_cell, "center")
        title_cell.paragraphs[0].clear()
        reset_cell_paragraph(title_cell.paragraphs[0], line_spacing=1.25)
        title_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_lines = [line.strip() for line in document.title.upper().split('\n') if line.strip()]
        for idx, line in enumerate(title_lines):
            if idx > 0:
                title_cell.paragraphs[0].add_run('\n')
            run = title_cell.paragraphs[0].add_run(line)
            format_run(run, TITLE_FONT_PT, bold=True)
            run.underline = True

        # Populate Logo Right
        merged_right = header_table.cell(0, 2)
        set_cell_width(merged_right, HEADER_LOGO_WIDTH_CM)
        set_cell_margins(merged_right, top_pt=2, left_pt=4, bottom_pt=2, right_pt=4)
        set_vertical_align(merged_right, "center")
        merged_right.paragraphs[0].clear()
        reset_cell_paragraph(merged_right.paragraphs[0])
        merged_right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if logo_right_bytes:
            run = merged_right.paragraphs[0].add_run()
            run.add_picture(BytesIO(logo_right_bytes), width=Cm(logo_right_dims[0]), height=Cm(logo_right_dims[1]))

        # Populate Cuadrante (optional label on line 1, value on line 2, centered)
        info_cell = header_table.cell(1, 1)
        set_cell_width(info_cell, HEADER_TITLE_WIDTH_CM)
        set_cell_margins(info_cell, top_pt=3, left_pt=5, bottom_pt=3, right_pt=5)
        set_vertical_align(info_cell, "center")
        info_cell.paragraphs[0].clear()
        reset_cell_paragraph(info_cell.paragraphs[0])
        info_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        resolved_label = (document.cuadrante_label or CUADRANTE_LABEL).strip()
        if document.show_cuadrante_label and resolved_label:
            label_run = info_cell.paragraphs[0].add_run(resolved_label)
            format_run(label_run, INFO_FONT_PT, bold=True)
            info_cell.paragraphs[0].add_run("\n")
        page_cuadrante = _breakable_text(_display_cuadrante(page.cuadrante or document.cuadrante))
        value_run = info_cell.paragraphs[0].add_run(page_cuadrante)
        format_run(value_run, INFO_FONT_PT, bold=True)

        # 2. Spacer = tabla 1x1 sin bordes (fiable en Word; line-spacing no lo es)
        spacer = doc.add_table(rows=1, cols=1)
        spacer.alignment = WD_TABLE_ALIGNMENT.CENTER
        spacer.autofit = False
        spacer.allow_autofit = False
        set_table_no_cell_margins(spacer)
        sp_pr = spacer._tbl.tblPr
        for tag in ("w:tblBorders",):
            existing = sp_pr.find(qn(tag))
            if existing is not None:
                sp_pr.remove(existing)
        sp_pr.append(parse_xml(_nil_table_borders_xml()))
        sp_pr.append(parse_xml(
            '<w:tblW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            f'w:w="{cm_to_twips(TABLE_WIDTH_CM)}" w:type="dxa"/>',
        ))
        sp_pr.append(parse_xml(
            '<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            'w:type="fixed"/>',
        ))
        spacer._tbl.tblGrid.gridCol_lst[0].set(qn("w:w"), str(cm_to_twips(TABLE_WIDTH_CM)))
        set_row_height(spacer.rows[0], GAP_UNDER_HEADER_CM)
        sp_cell = spacer.cell(0, 0)
        set_cell_width(sp_cell, TABLE_WIDTH_CM)
        set_cell_margins(sp_cell)
        sp_cell.paragraphs[0].clear()
        reset_cell_paragraph(sp_cell.paragraphs[0])
        tc_pr = sp_cell._tc.get_or_add_tcPr()
        for old in tc_pr.findall(qn("w:tcBorders")):
            tc_pr.remove(old)
        tc_pr.append(parse_xml(_nil_cell_borders_xml()))

        # 3. Photos: mismo grid que SheetPreview + marco exterior
        photos_table = doc.add_table(rows=PHOTO_TABLE_ROWS, cols=PHOTO_TABLE_COLS)
        photos_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        photos_table.autofit = False
        photos_table.allow_autofit = False
        set_table_no_cell_margins(photos_table)

        tblPr = photos_table._tbl.tblPr
        for tag in ("w:tblBorders",):
            existing = tblPr.find(qn(tag))
            if existing is not None:
                tblPr.remove(existing)
        tblPr.append(parse_xml(_photos_table_borders_xml()))
        tblPr.append(parse_xml(
            '<w:tblW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            f'w:w="{cm_to_twips(TABLE_WIDTH_CM)}" w:type="dxa"/>',
        ))
        tblPr.append(parse_xml(
            '<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            'w:type="fixed"/>',
        ))

        photo_col_widths: list[float] = [PHOTO_GAP_CM]
        for _ in range(PHOTO_COLS):
            photo_col_widths.append(PHOTO_WIDTH_CM)
            photo_col_widths.append(PHOTO_GAP_CM)
        grid = photos_table._tbl.tblGrid
        for col, width_cm in zip(grid.gridCol_lst, photo_col_widths, strict=True):
            col.set(qn("w:w"), str(cm_to_twips(width_cm)))

        set_row_height(photos_table.rows[0], PHOTO_GAP_CM)
        set_row_height(photos_table.rows[1], PHOTO_HEIGHT_CM)
        set_row_height(photos_table.rows[2], PHOTO_GAP_CM)
        set_row_height(photos_table.rows[3], PHOTO_HEIGHT_CM)
        set_row_height(photos_table.rows[4], PHOTO_GAP_CM)

        def _set_cell_frame(
            cell: Any,
            *,
            top: bool = False,
            bottom: bool = False,
            left: bool = False,
            right: bool = False,
        ) -> None:
            tcPr = cell._tc.get_or_add_tcPr()
            for old in tcPr.findall(qn("w:tcBorders")):
                tcPr.remove(old)
            tcPr.append(parse_xml(_cell_borders_xml(top=top, bottom=bottom, left=left, right=right)))

        def _style_empty_gap_cell(
            cell: Any,
            width_cm: float,
            *,
            top: bool = False,
            bottom: bool = False,
            left: bool = False,
            right: bool = False,
        ) -> None:
            set_cell_width(cell, width_cm)
            set_cell_margins(cell)
            set_vertical_align(cell, "center")
            cell.paragraphs[0].clear()
            reset_cell_paragraph(cell.paragraphs[0])
            run = cell.paragraphs[0].add_run("")
            format_run(run, 1)
            _set_cell_frame(cell, top=top, bottom=bottom, left=left, right=right)

        # Filas gap + marco exterior en perímetro (tcBorders; tblBorders solo no basta en Word)
        last_col = PHOTO_TABLE_COLS - 1
        for gap_row_idx in (0, 2, 4):
            gap_row = photos_table.rows[gap_row_idx]
            for col_idx, width_cm in enumerate(photo_col_widths):
                _style_empty_gap_cell(
                    gap_row.cells[col_idx],
                    width_cm,
                    top=gap_row_idx == 0,
                    bottom=gap_row_idx == 4,
                    left=col_idx == 0,
                    right=col_idx == last_col,
                )

        slot_map = {ref.position: ref.filename for ref in page.images if 1 <= ref.position <= MAX_SLOTS}
        for row_idx in range(PHOTO_ROWS):
            table_row_idx = 1 if row_idx == 0 else 3
            photo_row = photos_table.rows[table_row_idx]
            for col_idx in range(PHOTO_COLS):
                table_col_idx = col_idx * 2 + 1
                position = row_idx * PHOTO_COLS + col_idx + 1
                cell = photo_row.cells[table_col_idx]
                set_cell_width(cell, PHOTO_WIDTH_CM)
                set_cell_margins(cell)
                set_vertical_align(cell, "top")
                cell.paragraphs[0].clear()
                reset_cell_paragraph(cell.paragraphs[0])
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_cell_frame(cell)  # interiores sin borde
                slot_filename = slot_map.get(position)
                content = image_bytes.get(slot_filename or "")
                if slot_filename and content and _valid_image_bytes(content):
                    run = cell.paragraphs[0].add_run()
                    run.add_picture(
                        BytesIO(content),
                        width=Cm(PHOTO_WIDTH_CM),
                        height=Cm(PHOTO_HEIGHT_CM),
                    )
                else:
                    run = cell.paragraphs[0].add_run("Sin imagen")
                    format_run(run, 8)
                    run.italic = True
                    run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

            for gap_col_idx in range(0, PHOTO_TABLE_COLS, 2):
                _style_empty_gap_cell(
                    photo_row.cells[gap_col_idx],
                    PHOTO_GAP_CM,
                    left=gap_col_idx == 0,
                    right=gap_col_idx == last_col,
                )

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue(), _default_filename("docx")
