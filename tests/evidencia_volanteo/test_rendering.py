
from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn
from pypdf import PdfReader

from backend.core.evidencia_volanteo import (
    EvidenciaDocument,
    EvidenciaPage,
    ImageRef,
    RenderingError,
    render_docx,
    render_pdf,
    render_pdf_html,
)
from backend.core.evidencia_volanteo.layout import (
    BORDER_PT,
    EMPTY_CUADRANTE_PLACEHOLDER,
    PHOTO_GAP_CM,
    PHOTO_TABLE_COLS,
    PHOTO_TABLE_ROWS,
)
from backend.core.evidencia_volanteo.rendering import (
    _build_image_uris,
    _jinja_env,
    _prepare_logos,
    _serialize_pages,
    layout_context,
)
from tests.weasyprint_env import weasyprint_native_available

requires_weasyprint = pytest.mark.skipif(
    not weasyprint_native_available(),
    reason="WeasyPrint native libraries (GTK/Pango) unavailable on this host",
)


def _tiny_png() -> str:
    return (
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg=="
    )


def _make_document() -> EvidenciaDocument:
    return EvidenciaDocument(
        title="EVIDENCIAS FOTOGRAFICAS DEL VOLANTEO",
        cuadrante="AV EL SOL - DISTRITO CHORRILLOS",
        pages=(
            EvidenciaPage(images=(
                ImageRef(filename="img1.jpg", position=1),
                ImageRef(filename="img2.jpg", position=2),
            )),
        ),
    )


def test_render_pdf_empty_pages_raises() -> None:
    doc = EvidenciaDocument(title="T", cuadrante="C", pages=())
    with pytest.raises(RenderingError, match="No hay páginas"):
        render_pdf(doc, {}, {})


def test_render_pdf_rejects_over_max_pages(monkeypatch: pytest.MonkeyPatch) -> None:
    import backend.core.evidencia_volanteo.rendering as rendering_mod

    monkeypatch.setattr(rendering_mod, "MAX_PAGES", 2)
    doc = EvidenciaDocument(
        title="T",
        cuadrante="C",
        pages=tuple(EvidenciaPage(images=()) for _ in range(3)),
    )
    with pytest.raises(RenderingError, match="máximo de 2 páginas"):
        render_pdf(doc, {}, {})


def test_render_docx_rejects_over_max_pages(monkeypatch: pytest.MonkeyPatch) -> None:
    import backend.core.evidencia_volanteo.rendering as rendering_mod

    monkeypatch.setattr(rendering_mod, "MAX_PAGES", 2)
    doc = EvidenciaDocument(
        title="T",
        cuadrante="C",
        pages=tuple(EvidenciaPage(images=()) for _ in range(3)),
    )
    with pytest.raises(RenderingError, match="máximo de 2 páginas"):
        render_docx(doc, {}, {})


def test_render_docx_empty_pages_raises() -> None:
    doc = EvidenciaDocument(title="T", cuadrante="C", pages=())
    with pytest.raises(RenderingError, match="No hay páginas"):
        render_docx(doc, {}, {})


@requires_weasyprint
def test_render_pdf_success() -> None:
    doc = _make_document()
    images = {"img1.jpg": _tiny_png(), "img2.jpg": _tiny_png()}
    pdf_bytes, filename = render_pdf(doc, {}, images)
    assert filename.endswith(".pdf")
    reader = PdfReader(BytesIO(pdf_bytes))
    assert len(reader.pages) == 1


def test_render_docx_success() -> None:
    doc = _make_document()
    images = {"img1.jpg": _tiny_png(), "img2.jpg": _tiny_png()}
    docx_bytes, filename = render_docx(doc, {}, images)
    assert filename.endswith(".docx")
    assert len(docx_bytes) > 1000


@requires_weasyprint
def test_render_pdf_with_logos() -> None:
    doc = _make_document()
    logos = {"left": _tiny_png(), "right": _tiny_png()}
    pdf_bytes, _ = render_pdf(doc, logos, {})
    assert len(pdf_bytes) > 500


@requires_weasyprint
def test_render_pdf_per_page_cuadrante() -> None:
    doc = EvidenciaDocument(
        title="TEST",
        cuadrante="",
        pages=(
            EvidenciaPage(images=(), cuadrante="ZONA A"),
            EvidenciaPage(images=(), cuadrante="ZONA B"),
        ),
    )
    pdf_bytes, _ = render_pdf(doc, {}, {})
    reader = PdfReader(BytesIO(pdf_bytes))
    assert len(reader.pages) == 2


def test_docx_border_width_matches_preview() -> None:
    doc = _make_document()
    docx_bytes, _ = render_docx(doc, {}, {})
    document = Document(BytesIO(docx_bytes))
    expected_sz = str(int(BORDER_PT * 8))
    header_borders = document.tables[0]._tbl.tblPr.find(qn("w:tblBorders"))
    assert header_borders is not None
    top = header_borders.find(qn("w:top"))
    assert top is not None
    assert top.get(qn("w:sz")) == expected_sz
    assert len(document.tables) >= 3
    spacer_borders = document.tables[1]._tbl.tblPr.find(qn("w:tblBorders"))
    assert spacer_borders is not None
    assert spacer_borders.find(qn("w:top")).get(qn("w:val")) == "nil"
    photos_borders = document.tables[2]._tbl.tblPr.find(qn("w:tblBorders"))
    assert photos_borders is not None
    assert photos_borders.find(qn("w:top")).get(qn("w:sz")) == expected_sz
    assert photos_borders.find(qn("w:insideH")).get(qn("w:val")) == "nil"
    corner = document.tables[2].cell(0, 0)._tc.tcPr.find(qn("w:tcBorders"))
    assert corner is not None
    assert corner.find(qn("w:top")).get(qn("w:val")) == "single"
    assert corner.find(qn("w:left")).get(qn("w:val")) == "single"
    assert corner.find(qn("w:top")).get(qn("w:sz")) == expected_sz
    inner_photo = document.tables[2].cell(1, 1)._tc.tcPr.find(qn("w:tcBorders"))
    assert inner_photo is not None
    assert inner_photo.find(qn("w:top")).get(qn("w:val")) == "nil"
    assert inner_photo.find(qn("w:left")).get(qn("w:val")) == "nil"


def test_docx_info_row_grows_for_long_cuadrante() -> None:
    long_value = "QWSADD" + ("D" * 80)
    doc = EvidenciaDocument(
        title="TEST",
        cuadrante="",
        pages=(EvidenciaPage(images=(), cuadrante=long_value),),
    )
    docx_bytes, _ = render_docx(doc, {}, {})
    document = Document(BytesIO(docx_bytes))
    info_height = document.tables[0].rows[1]._tr.trPr.find(qn("w:trHeight"))
    assert info_height is not None
    assert info_height.get(qn("w:hRule")) == "atLeast"
    info_text = document.tables[0].cell(1, 1).paragraphs[0].text.replace("\u200b", "")
    assert long_value.upper() in info_text or long_value in info_text


def test_empty_cuadrante_shows_placeholder_in_docx() -> None:
    doc = EvidenciaDocument(
        title="TEST",
        cuadrante="",
        pages=(EvidenciaPage(images=(), cuadrante=""),),
    )
    docx_bytes, _ = render_docx(doc, {}, {})
    document = Document(BytesIO(docx_bytes))
    info_text = document.tables[0].cell(1, 1).paragraphs[0].text
    assert EMPTY_CUADRANTE_PLACEHOLDER in info_text


def test_docx_custom_cuadrante_label() -> None:
    doc = EvidenciaDocument(
        title="TEST",
        cuadrante="ZONA A",
        pages=(EvidenciaPage(images=(), cuadrante="ZONA A"),),
        cuadrante_label="SECTOR INTERVENIDO:",
        show_cuadrante_label=True,
    )
    docx_bytes, _ = render_docx(doc, {}, {})
    document = Document(BytesIO(docx_bytes))
    info_text = document.tables[0].cell(1, 1).paragraphs[0].text
    assert "SECTOR INTERVENIDO:" in info_text
    assert "CUADRANTE AFECTADO:" not in info_text
    assert "ZONA A" in info_text


def test_docx_hides_cuadrante_label_when_disabled() -> None:
    doc = EvidenciaDocument(
        title="TEST",
        cuadrante="ZONA B",
        pages=(EvidenciaPage(images=(), cuadrante="ZONA B"),),
        show_cuadrante_label=False,
    )
    docx_bytes, _ = render_docx(doc, {}, {})
    document = Document(BytesIO(docx_bytes))
    info_text = document.tables[0].cell(1, 1).paragraphs[0].text
    assert "CUADRANTE AFECTADO:" not in info_text
    assert "ZONA B" in info_text


def test_pdf_html_cuadrante_matches_preview() -> None:
    doc = EvidenciaDocument(
        title="TEST",
        cuadrante="",
        pages=(EvidenciaPage(images=(), cuadrante="zona norte"),),
    )
    template = _jinja_env.get_template("evidencia-volanteo.html")
    pages_data = _serialize_pages(doc, _build_image_uris({}, None))
    html = template.render({
        "title": doc.title,
        "cuadrante": doc.cuadrante,
        "pages": pages_data,
        "logo_left": _prepare_logos({})[0],
        "logo_right": _prepare_logos({})[1],
        **layout_context(),
    })
    assert "ZONA NORTE" in html
    assert "0.75pt solid #000" in html


def test_empty_cuadrante_shows_placeholder_in_pdf_html() -> None:
    doc = EvidenciaDocument(
        title="TEST",
        cuadrante="",
        pages=(EvidenciaPage(images=(), cuadrante=""),),
    )
    template = _jinja_env.get_template("evidencia-volanteo.html")
    pages_data = _serialize_pages(doc, _build_image_uris({}, None))
    html = template.render({
        "title": doc.title,
        "cuadrante": doc.cuadrante,
        "pages": pages_data,
        "logo_left": None,
        "logo_right": None,
        **layout_context(),
    })
    assert EMPTY_CUADRANTE_PLACEHOLDER in html


def test_pdf_html_custom_and_hidden_cuadrante_label() -> None:
    doc = EvidenciaDocument(
        title="TEST",
        cuadrante="",
        pages=(EvidenciaPage(images=(), cuadrante="zona norte"),),
        cuadrante_label="AREA AFECTADA:",
        show_cuadrante_label=True,
    )
    template = _jinja_env.get_template("evidencia-volanteo.html")
    pages_data = _serialize_pages(doc, _build_image_uris({}, None))
    html = template.render({
        "title": doc.title,
        "cuadrante": doc.cuadrante,
        "pages": pages_data,
        "logo_left": None,
        "logo_right": None,
        **layout_context(),
        "cuadrante_label": doc.cuadrante_label,
        "show_cuadrante_label": doc.show_cuadrante_label,
    })
    assert "AREA AFECTADA:" in html
    assert "CUADRANTE AFECTADO:" not in html

    html_hidden = template.render({
        "title": doc.title,
        "cuadrante": doc.cuadrante,
        "pages": pages_data,
        "logo_left": None,
        "logo_right": None,
        **layout_context(),
        "cuadrante_label": "CUADRANTE AFECTADO:",
        "show_cuadrante_label": False,
    })
    assert "CUADRANTE AFECTADO:" not in html_hidden
    assert "ZONA NORTE" in html_hidden


@requires_weasyprint
def test_render_pdf_html_from_preview_markup() -> None:
    html = """<!DOCTYPE html><html><head><meta charset='utf-8'></head>
    <body><div style='width:210mm;height:297mm'>Preview</div></body></html>"""
    pdf_bytes, filename = render_pdf_html(html)
    assert filename.endswith('.pdf')
    assert pdf_bytes.startswith(b'%PDF')


def test_build_image_uris_embeds_disk_paths_as_data_uri(tmp_path: Path) -> None:
    import base64
    from pathlib import Path

    png = base64.b64decode(_tiny_png(), validate=True)
    image_path = Path(tmp_path) / "foto.png"
    image_path.write_bytes(png)

    uris = _build_image_uris({}, {"foto.png": str(image_path)})
    assert "foto.png" in uris
    assert uris["foto.png"].startswith("data:image/png;base64,")
    assert "file:" not in uris["foto.png"]


@requires_weasyprint
def test_render_pdf_disk_backed_images_embed(tmp_path: Path) -> None:
    import base64
    from pathlib import Path

    png = base64.b64decode(_tiny_png(), validate=True)
    image_path = Path(tmp_path) / "img1.png"
    image_path.write_bytes(png)
    doc = EvidenciaDocument(
        title="TEST",
        cuadrante="C1",
        pages=(EvidenciaPage(images=(ImageRef(filename="img1.jpg", position=1),)),),
    )
    pdf_bytes, _ = render_pdf(
        doc,
        logos={},
        images={},
        image_paths={"img1.jpg": str(image_path)},
    )
    reader = PdfReader(BytesIO(pdf_bytes))
    assert any(page.images for page in reader.pages), "disk-backed image missing from PDF"


def test_build_image_uris_preserves_gif_mime(tmp_path: Path) -> None:
    from PIL import Image

    image_path = tmp_path / "foto.gif"
    image = Image.new("P", (1, 1))
    image.save(image_path, format="GIF")

    uris = _build_image_uris({}, {"foto.gif": str(image_path)})

    assert uris["foto.gif"].startswith("data:image/gif;base64,")


def test_docx_long_cuadrante_stays_editable_text() -> None:
    long_value = "QWSADD" + ("D" * 80)
    doc = EvidenciaDocument(
        title="TEST",
        cuadrante="",
        pages=(EvidenciaPage(images=(), cuadrante=long_value),),
    )
    docx_bytes, _ = render_docx(doc, {}, {})
    document = Document(BytesIO(docx_bytes))
    info_text = document.tables[0].cell(1, 1).paragraphs[0].text
    assert "QWSADD" in info_text
    assert info_text.replace("\u200b", "").endswith("D" * 10)
    assert len(document.tables) == 3


@requires_weasyprint
def test_render_pdf_html_six_images_one_a4_page() -> None:
    tiny = _tiny_png()
    data_uri = f"data:image/png;base64,{tiny}"
    row1 = "".join(
        f'<td style="border:0.75pt solid #000;height:12cm;overflow:hidden;padding:0">'
        f'<img src="{data_uri}" style="width:100%;height:12cm;object-fit:contain;display:block;'
        f'max-height:12cm"/></td>'
        for _ in range(3)
    )
    row2 = "".join(
        f'<td style="border:0.75pt solid #000;height:12cm;overflow:hidden;padding:0">'
        f'<img src="{data_uri}" style="width:100%;height:12cm;object-fit:contain;display:block;'
        f'max-height:12cm"/></td>'
        for _ in range(3)
    )
    html = f"""<!DOCTYPE html><html><head><meta charset='utf-8'><style>
    @page {{ size: A4 portrait; margin: 8mm; }}
    .ev-sheet-page {{
      width: 100%; height: 27.8cm; max-height: 27.8cm; overflow: hidden;
      page-break-inside: avoid; page-break-after: always;
    }}
    </style></head><body>
    <div class="ev-sheet-page"><table style="width:19.4cm;border-collapse:collapse;table-layout:fixed">
    <tr style="height:12cm">{row1}</tr>
    <tr style="height:0.6cm"><td colspan="3" style="border:0.75pt solid #000"></td></tr>
    <tr style="height:12cm">{row2}</tr>
    </table></div></body></html>"""
    pdf_bytes, _ = render_pdf_html(html)
    reader = PdfReader(BytesIO(pdf_bytes))
    assert len(reader.pages) == 1


def test_pdf_html_uses_contain_for_photos() -> None:
    doc = _make_document()
    template = _jinja_env.get_template("evidencia-volanteo.html")
    pages_data = _serialize_pages(doc, _build_image_uris({"img1.jpg": _tiny_png()}, None))
    html = template.render({
        "title": doc.title,
        "cuadrante": doc.cuadrante,
        "pages": pages_data,
        "logo_left": None,
        "logo_right": None,
        **layout_context(),
    })
    assert "object-fit: fill" in html
    assert "object-position: center" in html


def test_pdf_html_has_horizontal_photo_gaps() -> None:
    doc = _make_document()
    template = _jinja_env.get_template("evidencia-volanteo.html")
    pages_data = _serialize_pages(doc, _build_image_uris({}, None))
    html = template.render({
        "title": doc.title,
        "cuadrante": doc.cuadrante,
        "pages": pages_data,
        "logo_left": None,
        "logo_right": None,
        **layout_context(),
    })
    assert f'colspan="{PHOTO_TABLE_COLS}"' in html
    assert "col-gap" in html
    assert f"width: {PHOTO_GAP_CM}cm" in html or f"width: {PHOTO_GAP_CM}cm;" in html


def test_docx_photos_table_has_gap_columns() -> None:
    doc = _make_document()
    docx_bytes, _ = render_docx(doc, {}, {})
    document = Document(BytesIO(docx_bytes))
    photos_table = document.tables[2]
    assert len(photos_table.columns) == PHOTO_TABLE_COLS
    assert len(photos_table.rows) == PHOTO_TABLE_ROWS

    doc = _make_document()
    images = {"img1.jpg": _tiny_png(), "img2.jpg": _tiny_png()}
    docx_bytes, _ = render_docx(doc, {}, images)
    document = Document(BytesIO(docx_bytes))
    photos_table = document.tables[2]
    shapes = photos_table._element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline")
    assert len(shapes) >= 2


@requires_weasyprint
def test_render_pdf_six_images_per_page() -> None:
    images = {}
    refs = []
    for i in range(1, 7):
        name = f"img{i}.jpg"
        images[name] = _tiny_png()
        refs.append(ImageRef(filename=name, position=i))
    doc = EvidenciaDocument(
        title="TEST",
        cuadrante="ZONA",
        pages=(EvidenciaPage(images=tuple(refs)),),
    )
    pdf_bytes, _ = render_pdf(doc, {}, images)
    reader = PdfReader(BytesIO(pdf_bytes))
    assert len(reader.pages) == 1
